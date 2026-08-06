#type: ignore
"""Ingestion module for miscellaneous context documents (txt, docx, code files).

Scans data/docs/ and creates Document objects with structured extracted text.
Run with: uv run -m chat_cembrowski.data.doc_ingestion
"""

import logging
import uuid
from pathlib import Path
from typing import Optional

from .models import Document
from .serialization import save_document, load_documents_from_json

logger = logging.getLogger(__name__)

DOCS_DIR = Path(__file__).resolve().parents[3] / "data" / "docs"
DOC_JSON_DIR = Path(__file__).resolve().parents[3] / "data" / "doc_json"

TEXT_EXTENSIONS = {".txt", ".md"}
DOCX_EXTENSIONS = {".docx"}
CODE_EXTENSIONS = {
    ".py", ".js", ".ts", ".jsx", ".tsx",
    ".c", ".cpp", ".h", ".hpp",
    ".java", ".r", ".m", ".sh", ".bash",
    ".sql", ".yaml", ".yml", ".toml", ".ini", ".cfg",
}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | DOCX_EXTENSIONS | CODE_EXTENSIONS


def _extract_txt(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8", errors="replace")


def _extract_docx(file_path: Path) -> str:
    """Extract text from a .docx file as structured markdown.

    Headings → # / ## / ###, list paragraphs → - bullets, tables → markdown tables.
    This structure improves chunking by giving the splitter clear semantic boundaries.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(str(file_path))

    # Build element → object lookups for O(1) access during body traversal.
    para_by_elem = {p._element: p for p in doc.paragraphs}
    table_by_elem = {t._element: t for t in doc.tables}

    lines: list[str] = []

    for block in doc.element.body:
        tag = block.tag.split("}")[-1]  # strip XML namespace

        if tag == "p":
            para = para_by_elem.get(block)
            if para is None:
                continue

            text = para.text.strip()
            if not text:
                lines.append("")
                continue

            style = para.style.name if para.style else ""

            if style.startswith("Heading"):
                try:
                    level = int(style.split()[-1])
                except (ValueError, IndexError):
                    level = 4
                lines.append(f"{'#' * level} {text}")
            elif "List" in style:
                lines.append(f"- {text}")
            else:
                lines.append(text)

        elif tag == "tbl":
            tbl = table_by_elem.get(block)
            if tbl is None:
                continue

            lines.append("")
            for i, row in enumerate(tbl.rows):
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
            lines.append("")

    return "\n".join(lines)


def _extract_code(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8", errors="replace")


def ingest_local_docs(
    docs_dir: Path = DOCS_DIR,
    doc_json_dir: Path = DOC_JSON_DIR,
) -> list[Document]:
    """Create Document objects for files in docs_dir not yet registered in doc_json_dir.

    Supported types: .txt, .md, .docx, and common code extensions.
    Idempotent — skips files whose source_file is already in doc_json_dir.

    Args:
        docs_dir: Source directory (default: data/docs)
        doc_json_dir: JSON output directory (default: data/doc_json)

    Returns:
        List of newly created Document objects.
    """
    docs_dir.mkdir(parents=True, exist_ok=True)
    doc_json_dir.mkdir(parents=True, exist_ok=True)

    known_source_files = {d.source_file for d in load_documents_from_json(doc_json_dir)}

    new_docs: list[Document] = []

    for file_path in sorted(docs_dir.iterdir()):
        if not file_path.is_file():
            continue
        if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            logger.debug(f"Skipping unsupported file: {file_path.name}")
            continue
        if file_path.name in known_source_files:
            logger.info(f"Already registered, skipping: {file_path.name}")
            continue

        logger.info(f"Ingesting: {file_path.name}")
        ext = file_path.suffix.lower()

        try:
            if ext in DOCX_EXTENSIONS:
                text = _extract_docx(file_path)
                file_type = "docx"
            elif ext in TEXT_EXTENSIONS:
                text = _extract_txt(file_path)
                file_type = ext.lstrip(".")
            else:
                text = _extract_code(file_path)
                file_type = ext.lstrip(".")
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path.name}: {e}")
            continue

        doc = Document(
            id=str(uuid.uuid7()),
            title=file_path.stem,
            source_file=file_path.name,
            file_type=file_type,
            text=text,
            processed=False,
        )
        save_document(doc, doc_json_dir)
        new_docs.append(doc)
        logger.info(f"Created Document: '{doc.title}' ({doc.file_type}, {len(doc.text):,} chars)")

    logger.info(f"Document ingestion complete. New documents: {len(new_docs)}")
    return new_docs


if __name__ == "__main__":
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
    )
    docs = ingest_local_docs()
    for doc in docs:
        print(f"- {doc.title} ({doc.file_type}, {len(doc.text):,} chars)")
